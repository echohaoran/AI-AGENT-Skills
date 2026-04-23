"""
Word document handler with text extraction and redaction.
Supports python-docx for .docx files.
"""

import copy
from typing import Optional, Any

from scripts.handlers.base import BaseHandler


class WordHandler(BaseHandler):
    """Handles Word .docx file reading and writing via python-docx."""

    def read(self, file_path: str) -> Optional[str]:
        """Extract text from Word document (paragraphs + tables)."""
        try:
            from docx import Document
            doc = Document(file_path)
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = [
                        cell.text.strip()
                        for cell in row.cells
                        if cell.text.strip()
                    ]
                    if row_text:
                        parts.append(" | ".join(row_text))
            return "\n".join(parts)
        except Exception:
            return None

    def write(self, file_path: str, content: str, original_path: str) -> bool:
        """
        Write redacted content back to a new .docx file.
        Preserves original document structure; replaces paragraph text in-place.

        Args:
            file_path: Destination path for the redacted file.
            content: Full redacted text content.
            original_path: Path to original file (used as source for structure).

        Returns:
            True if successful, False otherwise.
        """
        try:
            from docx import Document

            doc = Document(original_path)

            # Rebuild paragraphs: split content by newlines, assign to doc paragraphs
            lines = content.split("\n")
            para_index = 0

            # Process paragraphs first
            for para in doc.paragraphs:
                if para.text.strip():
                    if para_index < len(lines):
                        para.text = lines[para_index]
                        para_index += 1

            # Process tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip() and para_index < len(lines):
                            cell.text = lines[para_index]
                            para_index += 1

            doc.save(file_path)
            return True
        except Exception as e:
            return False

    def write_with_mapping(
        self,
        file_path: str,
        original_path: str,
        replacements: dict[str, str],
    ) -> bool:
        """
        Write redacted file using exact keyword→placeholder mapping.
        Preserves document structure by doing in-place replacements.

        Args:
            file_path: Destination path for the redacted file.
            original_path: Path to original file.
            replacements: Dict of {original_keyword: placeholder_marker}.

        Returns:
            True if successful, False otherwise.
        """
        try:
            from docx import Document

            doc = Document(original_path)

            # Sort replacements by length (longest first) to prevent short matches
            # from consuming parts of longer keywords
            sorted_replacements = sorted(
                replacements.items(), key=lambda x: -len(x[0])
            )

            def apply_replacements(text: str) -> str:
                result = text
                for original, placeholder in sorted_replacements:
                    if original in result:
                        # placeholder can be "[████TYPE_1]" or "<<REDACTED_TYPE_1>>"
                        replacement = f"[{placeholder}]" if not placeholder.startswith("<<") else placeholder
                        result = result.replace(original, replacement)
                return result

            for para in doc.paragraphs:
                if para.text.strip():
                    para.text = apply_replacements(para.text)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            cell.text = apply_replacements(cell.text)

            doc.save(file_path)
            return True
        except Exception:
            return False
